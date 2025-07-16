import docx
import os
currentDir = os.path.dirname(os.path.abspath(__file__))
dataSourceDir = os.path.join(currentDir, "datasource")
 
doc1 = docx.Document(os.path.join(dataSourceDir, "word_doc_1.docx"))
doc2 = docx.Document(os.path.join(dataSourceDir, "word_doc_2.docx"))
doc1paragraphs = []
doc2paragraphs = []

for paragraph in doc1.paragraphs: #We save the paragraphs in lists
    doc1paragraphs.append(paragraph.text)
for paragraph in doc2.paragraphs:
    doc2paragraphs.append(paragraph.text)
 
for i in doc1paragraphs: #We check which paragraphs match and which do not
    if i in doc2paragraphs:
         print(f"[MATCH   ] {i}")
    else:
         print(f"[NO MATCH] {i}")